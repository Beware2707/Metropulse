"""The main MICE project proposal, rendered as both .docx and .pdf.

The content is defined once as a section list and rendered twice, for the same
reason the deck is: a proposal whose Word and PDF versions disagree is worse
than either on its own.
"""

from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Spacer

from . import theme as T
from .facts import APP_STACK, BACKEND_STACK, DATA_REQUIRED, INFRA, f
from .report import (CONTENT_W, bullets, build, callout, flow_diagram, heading,
                     para, status_row, table)

# Content is a list of (kind, payload) so both renderers walk the same tree.
# kinds: h1, h2, p, lead, ul, status, table, callout, flow, pagebreak


def _content() -> list[tuple]:
    C: list[tuple] = []
    a = C.append

    a(("h1", "MICE at a glance"))
    a(("table", ([
        ["", "The proposal in one page"],
        ["Problem",
         "Delhi Metro commuters must make multiple decisions during a "
         "journey — gate, platform, coach, interchange, exit, step-free "
         "route — that conventional route planners do not address."],
        ["Solution",
         "A journey companion providing contextual platform, gate, "
         "accessibility, journey and realtime guidance at the moment each "
         "decision is made."],
        ["Proof",
         "A working Android application and deployed backend · "
         f"{f('backend_tests')} backend tests · {f('app_tests')} application "
         f"tests · {f('beta_users')} commuter trial participants · "
         f"{f('lifts')} lifts, {f('gates')} gates and {f('platforms')} "
         f"platforms integrated · {f('pathway_stations')} OTD-mapped "
         "stations."],
        ["Gap", "Official Delhi Metro realtime and operational data."],
        ["MICE opportunity",
         "Integrate → validate → controlled pilot → measure passenger "
         "impact."],
        ["Ask",
         "Realtime data, operational data, aggregated passenger-flow data, "
         "technical collaboration, and a controlled ninety-day pilot."],
    ], [32 * mm, CONTENT_W - 32 * mm])))
    a(("callout", ("No funding or exclusivity is requested at this stage",
                   "This page is the whole proposal in one screen: a working "
                   "product on approved data, one missing input, and a "
                   "measurable pilot to close the gap.", "good")))
    a(("pagebreak", None))

    a(("h1", "1. Executive summary"))
    a(("lead",
       "A journey is eight decisions, not one route — and route planning "
       "answers only the first. MetroPulse is the layer that answers the "
       "other seven: a working Android application and backend that turns "
       "Delhi Open Transit Data into moment-by-moment commuter guidance. It is "
       "deployed, tested, and has been used by commuters on the Delhi Metro. "
       "It does not have official realtime Delhi Metro data, and it says so "
       "everywhere that matters — in the product, in this proposal, and on the "
       "first slide of the accompanying deck."))
    a(("p",
       "The submission asks DMRC for five things: realtime Metro data, "
       "operational station data, aggregated passenger-flow data, technical "
       "collaboration on feed formats and identifiers, and a controlled "
       "ninety-day pilot. It does not ask for funding, exclusivity, or any "
       "operational role."))
    a(("p",
       "The architectural point most relevant to DMRC is that the realtime "
       "pipeline is already built and only its first block is missing. "
       "Connecting an official feed is an adapter and configuration task, held "
       "to a conformance suite that already exists in the repository."))
    a(("h2", "Status at a glance"))
    a(("table", ([
        ["Area", "Position today"],
        ["Application", f"Built and deployed. {f('app_tests')} tests passing."],
        ["Backend", f"Built and deployed on AWS. {f('backend_tests')} tests passing, "
                    f"{f('endpoints')} endpoints."],
        ["Official data", f"{f('lifts')} lifts, {f('gates')} gates, "
                          f"{f('platforms')} platforms integrated from OTD."],
        ["Realtime Metro data", "Not available. Pipeline built and waiting."],
        ["Crowd measurement", "Not available. Falls back to a generic prior."],
        ["Escalator data", "Does not exist in the approved datasets."],
        ["Commuter trial", f"{f('beta_users')} riders; defect-driven, not instrumented."],
        ["Organisation", f"{f('incorporation')}. {f('dpiit')}."],
    ], [42 * mm, CONTENT_W - 42 * mm])))
    a(("pagebreak", None))

    a(("h1", "2. How to read this proposal"))
    a(("p", "Every capability in this document carries exactly one label."))
    a(("table", ([["Label", "Meaning"]] +
                 [[k, v] for k, v in T.STATUS_MEANING.items()],
                 [45 * mm, CONTENT_W - 45 * mm])))
    a(("callout", ("On numbers",
                   "Every figure here is either measured, or labelled on the "
                   "page as reported rather than measured. Where a figure "
                   "would ordinarily be expected but neither is true, it is "
                   "shown as [DATA REQUIRED] rather than estimated.", "info")))

    a(("h1", "3. The commuter problem"))
    a(("p",
       "A metro journey is not one decision. It is a sequence of them, each "
       "with a different answer at a different moment: when to leave, which "
       "gate to enter, which platform, which coach, where to change, how "
       "crowded it will be, which exit to take, and — for a significant "
       "minority of riders — whether there is a step-free path at all."))
    a(("p",
       "A route planner answers the first of these and stops. The remaining "
       "decisions are made on concourses and platforms, usually under time "
       "pressure, from signage that assumes the rider already knows the "
       "station. The cost of getting them wrong is small individually and "
       "large in aggregate: a wrong exit adds a surface walk, a wrong coach "
       "adds four minutes, a wrong assumption about lifts can end a journey."))

    a(("h1", "4. What MetroPulse does"))
    a(("flow", ([("Plan", "Route and\npreference"), ("Leave", "When to\nleave"),
                 ("Enter", "Which\ngate"), ("Platform", "Which\nplatform"),
                 ("Board", "Which\ncoach"), ("Interchange", "Where to\nchange"),
                 ("Crowd", "How\nbusy"), ("Exit", "Which\nexit")], 4)))
    a(("status", ("IMPLEMENTED",
                  "Plan, leave, enter, platform, board, interchange and exit "
                  "guidance are built and in the shipped application.")))
    a(("status", ("REQUIRES DMRC DATA",
                  "Crowd guidance, which needs operational occupancy data to be "
                  "a measurement rather than a model.")))
    a(("p",
       "The application's flagship screen advances through these moments and "
       "shows only the one in front of the rider. The design principle is "
       "subtraction: a rider walking through a station with a bag in one hand "
       "can act on one instruction, not on a dashboard."))
    a(("pagebreak", None))

    a(("h1", "5. Current implementation"))
    a(("h2", "Backend"))
    a(("ul", BACKEND_STACK))
    a(("h2", "Application"))
    a(("ul", APP_STACK))
    a(("h2", "Deployment"))
    a(("ul", INFRA))
    a(("h2", "Scale"))
    a(("table", ([
        ["Measure", "Value"],
        ["Backend modules", f("py_modules")],
        ["Application source files", f("dart_files")],
        ["REST endpoints", f("endpoints")],
        ["Database migrations", f("migrations")],
        ["Backend tests", f("backend_tests")],
        ["Application tests", f("app_tests")],
        ["Release APK", f("apk_size")],
    ], [70 * mm, CONTENT_W - 70 * mm])))

    a(("h1", "6. Official data integration"))
    a(("status", ("APPROVED DATA", "All counts below are read from the loaded "
                                   "artifacts, not estimated.")))
    a(("table", ([
        ["Artifact", "Contents", "Count"],
        ["pathways.json", "Lifts", f("lifts")],
        ["pathways.json", "Entry / exit gates", f("gates")],
        ["pathways.json", "Platforms", f("platforms")],
        ["pathways.json", "Stations mapped", f("pathway_stations")],
        ["official_stations_gates.json", "Station and gate registry", f("gate_registry")],
        ["hourly_profile.json", "Hourly load entries", f("hourly_profile")],
        ["od_top_destinations.json", "Origin stations", f("od_origins")],
    ], [58 * mm, CONTENT_W - 58 * mm - 22 * mm, 22 * mm])))
    a(("callout", ("Escalators are not in the approved data",
                   "All five normalized OTD artifacts contain zero escalator "
                   "records. MetroPulse therefore makes no escalator claim "
                   "anywhere in the product, and none in this proposal. "
                   "Escalator inventory is requested from DMRC instead.", "warn")))
    a(("callout", ("The timetable has no Sunday service",
                   f"The published GTFS contains {f('trips_weekday')} weekday "
                   f"trips, {f('trips_saturday')} Saturday trips and "
                   f"{f('trips_sunday')} Sunday trips. The calendar declares a "
                   "Sunday service but no trip references it. The application "
                   "detects this and tells the rider it has no timetable for "
                   "the day rather than implying no trains are running.", "warn")))

    a(("h1", "7. Realtime architecture"))
    a(("flow", ([("DMRC realtime", "Requested"), ("Adapter", "Built"),
                 ("Validation", "Built"), ("ID mapping", "Built"),
                 ("Normalized position", "Built"), ("Redis", "Built"),
                 ("WebSocket", "Built"), ("Application", "Built")], 4)))
    a(("p",
       "Positions enter through a single interface and everything downstream "
       "consumes one normalized model carrying an explicit source label. That "
       "label is what allows the application to display SCHEDULE rather than "
       "LIVE over interpolated data — and what would make official DMRC data "
       "visibly distinct the moment it arrived."))
    a(("h2", "Today, and with a DMRC feed"))
    a(("p",
       "The same screens answer differently once real positions arrive. "
       "Nothing on the right needs a new application; each row is a different "
       "value behind the same interface, carrying a different provenance "
       "label. One distinction matters: with a feed, the train position "
       "becomes LIVE, but a computed arrival time is still labelled "
       "PREDICTED — a calculation never borrows its input's label."))
    a(("table", ([
        ["", "Today — schedule only", "With a DMRC realtime feed"],
        ["Train position", "Interpolated from the timetable",
         "Actual position from the feed — labelled LIVE"],
        ["Arrival time", "Scheduled, labelled SCHEDULE",
         "Predicted from LIVE vehicle position — labelled PREDICTED"],
        ["Disruption", "Rider reports only",
         "Detected from the feed within seconds"],
        ["Journey tracking", "GPS above ground, stop-count below",
         "Confirmed against the train the rider is on"],
        ["Crowd guidance", "A generic prior, the same everywhere",
         "Measured loading, per line and hour"],
        ["Interchange advice", "Static walking allowance",
         "Timed against the connecting train"],
    ], [34 * mm, (CONTENT_W - 34 * mm) / 2, (CONTENT_W - 34 * mm) / 2])))

    a(("h2", "Four kinds of data, kept separate"))
    a(("table", ([
        ["Kind", "Status today", "Shown to the rider as"],
        ["Official DMRC realtime", "Not available",
         "LIVE — reserved, never shown today"],
        ["MetroPulse estimate", "In use", "SCHEDULE"],
        ["AI prediction", "In use for delay and coach", "Labelled as estimated"],
        ["Scheduled data", "In use", "Timetable"],
    ], [42 * mm, CONTENT_W - 42 * mm - 42 * mm, 42 * mm])))
    a(("callout", ("The public OTD VehiclePositions feed is not Metro data",
                   "Checked against the live endpoint: vehicle identifiers are "
                   "road registration plates, several thousand vehicles report "
                   "simultaneously, and positions fall across ordinary roads. "
                   "It is Delhi's citywide bus GPS. MetroPulse ships with "
                   "realtime ingestion disabled for this reason. The full measurement "
                   "record is in the appendix.", "warn")))

    a(("h1", "8. AI and voice"))
    a(("status", ("IMPLEMENTED",
                  "On-device voice assistant with a deterministic intent "
                  "classifier: route, when to leave, which coach, fare, next "
                  "station, running late. Offline station search across names, "
                  "aliases and landmarks.")))
    a(("status", ("IMPLEMENTED",
                  "Coach recommendation ranking expected crowding against exit "
                  "alignment; delay and commute prediction behind pluggable "
                  "service interfaces.")))
    a(("status", ("REQUIRES DMRC DATA",
                  "Crowd prediction currently falls back to a generic prior "
                  "because no occupancy observations exist to learn from.")))
    a(("p",
       "No large language model is used. The assistant is a deterministic "
       "classifier over the application's own data, so every answer is a real "
       "figure from the same source the rest of the product shows. It answers "
       "only Delhi Metro questions and declines everything else by design."))

    a(("h1", "9. Accessibility"))
    a(("p",
       "Step-free guidance is built on DMRC's own pathway graph, and every "
       "claim is pitched at the evidence behind it. A gate is described as "
       "step-free only when the graph connects it to a lift-served platform."))
    a(("table", ([
        ["Tier", "Meaning", "Wording used"],
        ["Confirmed", "Graph connects this gate to a lift-served platform",
         "Lift to the platform"],
        ["Not mapped", "The station has lifts, but no mapped path from this gate",
         "Path from this gate not mapped"],
        ["No data", "Nothing is known", "Nothing is said"],
    ], [28 * mm, CONTENT_W - 28 * mm - 48 * mm, 48 * mm])))
    a(("p",
       f"Coverage is partial and stated as such: {f('acc_records')} of "
       f"{f('stations')} stations have an accessibility record, and "
       f"{f('acc_stepfree')} have at least one confirmed step-free gate. The "
       "application never presents 'unmapped' as 'inaccessible' — the "
       "distinction matters to the rider it affects most."))

    a(("h1", "10. Crowd management"))
    a(("flow", ([("Approved data", "Load profile,\nOD pairs"),
                 ("Realtime + historical", "Occupancy\nsignals"),
                 ("Crowd intelligence", "Aggregation with\nprovenance"),
                 ("Detection / prediction", "Where and when\nit builds"),
                 ("Passenger guidance", "One instruction"),
                 ("Rider acts", "Gate, interchange\nor departure time")], 3)))
    a(("callout", ("Scope boundary",
                   "MetroPulse complements DMRC's operational crowd-management "
                   "systems and does not replace them. Control-room detection, "
                   "staff deployment and station regulation remain DMRC's. "
                   "MetroPulse's role is the passenger-facing half: one clear "
                   "instruction before a rider reaches a congested concourse.",
                   "good")))
    a(("pagebreak", None))

    a(("h1", "Data trust model"))
    a(("lead",
       "Five levels of certainty, ordered. A lower level never silently "
       "overwrites a higher one, and nothing is displayed without its label — "
       "which is what lets a rider tell an official fact from an estimate "
       "without reading documentation."))
    a(("table", ([
        ["Level", "Source", "How the rider sees it", "Status"],
        ["1", "Official DMRC information",
         "Presented as official, attributed to DMRC", "REQUIRES DMRC DATA"],
        ["2", "Official OTD / static data",
         "Timetable, pathways and gates, shown as published data",
         "APPROVED DATA"],
        ["3", "MetroPulse calculation",
         "Routing, fares and interchange timing, marked SCHEDULE where "
         "time-based", "IMPLEMENTED"],
        ["4", "AI prediction",
         "Labelled as an estimate, never as fact", "IMPLEMENTED"],
        ["5", "Crowdsourced observation",
         "Attributed to riders, and only after independent confirmation",
         "IMPLEMENTED"],
    ], [14 * mm, 44 * mm, CONTENT_W - 14 * mm - 44 * mm - 34 * mm, 34 * mm])))
    a(("p",
       "The rule is enforced in the code rather than in review. A position "
       "carries its source label through every layer to the screen, so the "
       "interface cannot render a schedule estimate as though it were live "
       "without the label coming with it. Rider observations are the only "
       "level that requires agreement from multiple independent people before "
       "anyone sees them, because it is the only level where one person can "
       "be wrong on purpose."))
    a(("callout", ("Why this is a design constraint and not a promise",
                   "Discipline decays. Where it mattered most, the API was "
                   "shaped so the unwanted thing cannot be expressed at all: "
                   "the analytics call has no parameter for a search query, "
                   "and there is no way to publish a position without its "
                   "source. A rule that cannot be broken by accident does not "
                   "need to be remembered.", "good")))
    a(("pagebreak", None))

    a(("h1", "Security and privacy"))
    a(("h2", "Privacy by design"))
    a(("flow", ([
        ("Rider location", "stays on the device\nby default"),
        ("Opt-in", "per journey, stoppable\nat any time"),
        ("Minimum processing", "only what the next\ninstruction needs"),
        ("Journey intelligence", "the product of the data,\nnot a copy of it"),
        ("Aggregation", "rider reports used only\nafter independent agreement"),
        ("No individual tracking", "shared trips expire; analytics\ncarries no station pair"),
    ], 3)))
    a(("p",
       "Official data, calculated data and rider-generated observations are "
       "kept in separate stores and never merged, so one can never be "
       "presented as another."))
    a(("ul", [
        "Riders are identified only by an anonymous device identifier — never "
        "a name, email or phone number.",
        "Location stays on the device except when a rider explicitly shares a "
        "trip, which they start and stop themselves.",
        "An opt-in analytics pipeline is built and switched off. It is scoped "
        "so it cannot carry a search query, a spoken phrase, or the stations a "
        "rider travelled between.",
        "Rider contributions to station data are opt-in, per-prompt, and "
        "require agreement from multiple independent riders before use.",
        "A published privacy policy describes exactly what is collected; it is "
        "maintained in the repository alongside the code it describes.",
    ]))
    a(("p",
       "Any DMRC data provided would be used only for commuter-facing guidance "
       "within the agreed pilot scope, labelled as official, and never "
       "presented as more certain than it is."))

    a(("h2", "Security"))
    a(("ul", [
        "The licensed Delhi Transport Stack API key is held server-side and "
        "proxied. It is not in the application package and not in version "
        "control, so a decompiled APK does not yield a working credential.",
        "Database and cache are reachable only from the application host; "
        "public access is closed at the security-group level.",
        "Schema changes go through versioned migrations applied in "
        "production, so the deployed schema is always a known revision.",
        "Loaders replace each dataset wholesale inside one transaction, so a "
        "failed re-run cannot leave a station described half by old data and "
        "half by new.",
        "Rate limiting and a TLS certificate on a proper domain are named as "
        "hardening still to do before any public pilot, and are listed as "
        "such in the risk register rather than implied to be finished.",
    ]))
    a(("h2", "Auditability"))
    a(("p",
       "Every rider-visible claim can be traced back to the artifact it came "
       "from: the loaders record which file and which run produced each row, "
       "and provenance labels survive to the screen. Official and "
       "user-generated data are stored in separate tables and never merged, "
       "so a rider observation cannot be mistaken for an official record even "
       "by a future query written by someone who was not there."))

    a(("h1", "12. Testing"))
    a(("status", ("IMPLEMENTED",
                  f"{f('backend_tests')} backend tests and {f('app_tests')} "
                  "application tests, all passing.")))
    a(("status", ("IMPLEMENTED",
                  f"A commuter trial with {f('beta_users')} riders on the Delhi "
                  "Metro. Every defect they reported was diagnosed, fixed and "
                  "verified; the fix record is traceable in the repository.")))
    a(("p",
       "The trial was defect-driven rather than instrumented. It surfaced "
       "problems that no test suite would have found — station names nobody "
       "types the way the feed spells them, an assistant that answered a "
       "question without doing anything, and a build configuration that "
       "pointed release builds at a developer machine. All are fixed."))
    a(("table", ([
        ["Metric", "Status"],
        ["Trial participants", f("beta_users") + " (reported, not app-measured)"],
        ["Journeys completed", f("journeys_completed")],
        ["Retention", f("retention")],
        ["Crash-free rate", f("crash_free_rate")],
    ], [60 * mm, CONTENT_W - 60 * mm])))
    a(("p",
       "The three [DATA REQUIRED] rows need the opt-in analytics build in "
       "riders' hands, which in turn needs the privacy policy amendment "
       "published. They are not estimated here."))

    a(("h1", "Proposed pilot"))
    a(("status", ("PROPOSED", "Every stage below is proposed, not agreed. "
                              "The start date and the pilot lines are DMRC "
                              "to choose.")))
    a(("table", ([
        ["Stage", "Activity", "Exit criterion"],
        ["Days 0-15 — Technical integration",
         "Terms agreed; adapter connected to a DMRC feed in a non-public "
         "environment; identifiers mapped; conformance suite run against the "
         "live feed.",
         "Feed successfully consumed and mapped."],
        ["Days 16-30 — Data validation",
         "Freshness, completeness and accuracy measured against ground truth "
         "on selected corridors, including underground sections.",
         "Agreed freshness and completeness thresholds achieved."],
        ["Days 31-60 — Controlled commuter pilot",
         "A predefined cohort rides selected routes with opt-in analytics "
         "enabled, measuring the metrics currently marked [DATA REQUIRED].",
         "The cohort completes its test journeys."],
        ["Days 61-75 — Crowd and accessibility experiments",
         "Selected crowd-guidance and step-free scenarios run and measured.",
         "Selected scenarios validated."],
        ["Days 76-90 — Evaluation",
         "Joint review written up including the negative findings.",
         "MICE receives the final performance report."],
    ], [36 * mm, CONTENT_W - 36 * mm - 40 * mm, 40 * mm])))
    a(("table", ([
        ["Outcome", "Meaning"],
        ["GO", "Extend beyond the pilot lines, on terms discussed then."],
        ["ITERATE", "Revise and re-run the stage that missed its criterion."],
        ["STOP", "Wind down. The findings are delivered either way."],
    ], [26 * mm, CONTENT_W - 26 * mm])))
    a(("p",
       "All three endings are defined in advance, so stopping is a normal "
       "outcome rather than an admission — the only arrangement under which "
       "a negative result gets reported honestly."))

    a(("h1", "Key performance indicators"))
    a(("p",
       "None of these is measured today. Baselines are set jointly in the "
       "first week of the commuter-pilot stage, so any improvement is "
       "measured against something both sides agreed to. Targets are "
       "deliberately not asserted, because a target set without a baseline is "
       "a guess."))
    a(("h2", "Headline: passenger decision success rate"))
    a(("p",
       "Did MetroPulse give the correct actionable instruction at the moment "
       "the passenger needed it? Correct gate, platform, interchange, "
       "accessibility route, exit and realtime status — each confirmed "
       "in-app at the moment of use. This is the product proposition made "
       "measurable."))
    a(("h2", "Passenger"))
    a(("table", ([
        ["KPI", "How it is measured", "Baseline"],
        ["Passenger decision success rate",
         "In-app confirmation at the moment of use, across the six decision "
         "kinds above", DATA_REQUIRED],
        ["Journey completion rate",
         "Opt-in analytics: journeys started against journeys reaching the "
         "destination station", DATA_REQUIRED],
        ["Correct-exit rate",
         "In-app confirmation prompt at journey end", DATA_REQUIRED],
        ["Step-free journeys completed",
         "Accessibility-mode sessions finished without falling back",
         DATA_REQUIRED],
        ["Rider-reported usefulness",
         "Short in-app survey at journey end", DATA_REQUIRED],
    ], [44 * mm, CONTENT_W - 44 * mm - 30 * mm, 30 * mm])))
    a(("h2", "System"))
    a(("table", ([
        ["KPI", "How it is measured", "Baseline"],
        ["Arrival-time error against ground truth",
         "Seconds, per line and per hour, against DMRC own record",
         DATA_REQUIRED],
        ["Position accuracy underground",
         "Metres, against station arrival events rather than a satellite fix "
         "that does not exist below ground", DATA_REQUIRED],
        ["Feed availability and staleness",
         "Adapter conformance metrics already emitted today", DATA_REQUIRED],
        ["Crash-free session rate",
         "Crash reporting, wired but not yet configured", DATA_REQUIRED],
    ], [44 * mm, CONTENT_W - 44 * mm - 30 * mm, 30 * mm])))
    a(("h2", "Crowd"))
    a(("table", ([
        ["KPI", "How it is measured", "Baseline"],
        ["Congestion detected against congestion observed",
         "MetroPulse detections against DMRC station reports for the same "
         "period", DATA_REQUIRED],
        ["Guidance acted on",
         "Share of warnings after which the rider chose an alternative gate, "
         "interchange or departure time", DATA_REQUIRED],
        ["Load shifted off a peak",
         "Distribution of chosen departure times before and after guidance",
         DATA_REQUIRED],
        ["False-positive warnings",
         "Reviewed jointly with DMRC; a warning nobody needed is a defect",
         DATA_REQUIRED],
    ], [44 * mm, CONTENT_W - 44 * mm - 30 * mm, 30 * mm])))

    a(("h1", "15. Value to DMRC"))
    a(("ul", [
        "<b>Passenger experience</b> — guidance at the moment of each decision, "
        "in the rider's pocket rather than on a website.",
        "<b>Accessibility</b> — step-free guidance built on DMRC's own pathway "
        "data, with evidence tiers that never overstate what is known.",
        "<b>Crowd management</b> — passenger-facing guidance that complements "
        "control-room systems.",
        "<b>Digital innovation</b> — a working demonstration of what open data "
        "enables when it is taken seriously.",
        "<b>Data utilisation</b> — approved OTD datasets turned into daily "
        "commuter value rather than sitting unused.",
    ]))

    a(("h1", "16. Business model"))
    a(("p",
       "MetroPulse is not incorporated and holds no startup or DPIIT "
       "registration. There is no revenue, no funding and no commercial "
       "arrangement of any kind. This submission does not request funding."))
    a(("p",
       "The founder's stated intent is that MetroPulse remains free for "
       "commuters. Any future sustainability model would be discussed with "
       "DMRC rather than assumed here."))

    a(("h1", "17. Scalability"))
    a(("ul", [
        "The API and worker are separate processes: API replicas scale "
        "horizontally behind a load balancer, with exactly one worker.",
        "Redis pub/sub delivers realtime diffs to every replica, so adding "
        "capacity does not change the data path.",
        "Only changed vehicle positions are stored and broadcast, so cost "
        "scales with movement rather than with fleet size.",
        "The application is offline-first: station data, the network map and "
        "search work without connectivity.",
        "The codebase is cross-platform; an iOS build is a packaging exercise "
        "rather than a rewrite.",
    ]))

    a(("h1", "18. Roadmap"))
    a(("status", ("IN DEVELOPMENT",
                  "Background journey tracking with a next-station notification; "
                  "opt-in rider contributions to station data.")))
    a(("status", ("PROPOSED",
                  "Production hardening — TLS and a domain, restricted firewall, "
                  "secrets management, rate limiting. An instrumented commuter "
                  "trial to produce the metrics currently marked "
                  "[DATA REQUIRED].")))
    a(("status", ("REQUIRES DMRC DATA",
                  "Live train tracking, measured crowding, escalator guidance "
                  "and weekend service guidance.")))
    a(("status", ("PROPOSED",
                  "Longer term: multimodal journeys combining Metro with bus "
                  "and last-mile, and an iOS application.")))

    a(("h1", "19. Requested collaboration"))
    a(("table", ([
        ["Priority", "Request"],
        ["1 · Realtime Metro data",
         "Vehicle positions or train movement, ETA or trip updates, and "
         "service alerts — in a test environment first."],
        ["2 · Operational station data",
         "Facility operational status, gate closures, accessibility changes, "
         "and escalator inventory and status."],
        ["3 · Aggregated passenger-flow data",
         "Station occupancy, gate throughput or other aggregated, "
         "non-personal crowd indicators — whichever DMRC can share."],
        ["4 · Technical collaboration",
         "Feed specification, identifier mapping, refresh expectations, and "
         "a named pilot technical contact."],
        ["5 · Pilot",
         "A controlled ninety-day pilot on selected stations and routes, "
         "evaluated against the KPIs in this proposal."],
    ], [44 * mm, CONTENT_W - 44 * mm])))
    a(("p", "No funding or exclusivity is requested at this stage."))

    a(("h1", "20. Risks and mitigations"))
    a(("table", ([
        ["Risk", "Mitigation"],
        ["Realtime identifiers do not match static GTFS",
         "A configurable mapping chain is already implemented: exact match, "
         "explicit map, then rewrite rules."],
        ["Positioning is poor underground",
         "Multi-source tracking is designed and partly built: satellite fix, "
         "coarse network fix, motion-based stop counting and timetable "
         "fallback — each labelled distinctly."],
        ["Crowd data is unavailable during the pilot",
         "Crowd guidance is a separable phase and can be deferred without "
         "affecting the rest."],
        ["The product proves not useful",
         "Phase 04 explicitly permits stopping; every phase has an exit check."],
        ["Single-developer capacity",
         "Acknowledged. The scope of each pilot phase is deliberately small."],
    ], [52 * mm, CONTENT_W - 52 * mm])))

    a(("pagebreak", None))
    a(("h1", "21. Current limitations"))
    a(("p", "Stated plainly, because a reviewer will find them anyway."))
    a(("ul", [
        "No official realtime Delhi Metro data. Train positions are "
        "interpolated from the timetable and labelled as estimates.",
        "No measured crowding. The crowd model falls back to a generic prior.",
        "No escalator data anywhere in the approved datasets.",
        "No weekend timetable: the feed contains no Sunday trips.",
        f"Partial accessibility coverage — {f('acc_stepfree')} of "
        f"{f('stations')} stations have a confirmed step-free gate.",
        "The backend is served over plain HTTP against an IP address; TLS and "
        "a domain are outstanding and are a release blocker.",
        "No instrumented commuter metrics, as described in section 12.",
        "Not incorporated; no company documents, financials or certificates "
        "exist and none are included.",
        "Background journey tracking is built but not yet verified on a device, "
        "and its detection thresholds have not been validated on a real train.",
    ]))

    a(("h1", "22. Appendices"))
    a(("h2", "A. Accompanying documents"))
    a(("ul", [
        "Pitch_Deck.pptx / .pdf — twenty slides",
        "Technical_Architecture.pdf — system design and the realtime seam",
        "OTD_Data_Integration.pdf — datasets loaded, and what they omit",
        "Commuter_Validation.pdf — automated coverage and the rider trial",
        "Crowd_Management.pdf — data to one passenger instruction",
        "Realtime_Data_Request.pdf — what is asked of DMRC, and the evidence",
        "Pilot_Proposal.pdf — ninety days, five stages, and the KPIs",
        "Risk_and_Mitigation.pdf — every foreseeable risk and what handles it",
        "Founder_Profile.pdf — self-reported background, kept separate",
        "Roadmap.pdf — ordered by dependency, not by date",
        "Document_Status.pdf — index and honest completion status",
        "Demo/ — the release APK, screenshots, and a walkthrough",
    ]))
    a(("h2", "B. Contact"))
    a(("ul", [f("founder"), f("contact_phone"), f("contact_email")]))
    a(("h2", "C. Information still to be supplied"))
    a(("ul", [
        f"Journey completion, retention and crash-free rate — {DATA_REQUIRED}. "
        "All three need the consent-gated analytics build in riders' hands.",
    ]))
    a(("h2", "D. Public VehiclePositions feed — measurement record"))
    a(("p",
       "A capture of the public OTD VehiclePositions endpoint (2026-07-05) "
       "was analysed rather than assumed. This record is why MetroPulse "
       "ships with realtime ingestion disabled."))
    a(("table", ([
        ["Measurement", "Result"],
        ["Vehicles reporting", f("feed_vehicles") + " in " + f("feed_window")],
        ["Identifiers matching a road number-plate pattern",
         f("feed_plate_ids")],
        ["Vehicles resolving to a Metro route", f("feed_resolved")],
        ["Median distance to the nearest Metro station",
         f("feed_median_dist")],
        ["Within 50 m of any Metro station", f("feed_near_station")],
        ["Furthest vehicle from the network", f("feed_max_dist")],
        ["Distinct route identifiers (the Metro GTFS has 36)",
         f("feed_routes")],
    ], [CONTENT_W - 40 * mm, 40 * mm])))
    _renumber(C)
    return C


def _renumber(content: list[tuple]) -> None:
    """Number the h1 sections in order.

    The numbers used to be written into each heading by hand, so inserting a
    section meant editing every heading after it — and the appendix
    cross-references drifted the first time that was missed.
    """
    n = 0
    for i, (kind, payload) in enumerate(content):
        if kind != "h1" or not isinstance(payload, str):
            continue
        title = re.sub(r"^\d+\.\s*", "", payload)
        if title.startswith("Appendices"):
            content[i] = (kind, title)
            continue
        n += 1
        content[i] = (kind, f"{n}. {title}")


# ------------------------------------------------------------------- docx
def _docx(path: str) -> str:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("MetroPulse")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*T.rgb(T.INK))
    s = doc.add_paragraph()
    r = s.add_run("MICE Project Proposal\nCommuter intelligence for the Delhi Metro")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(*T.rgb(T.PRIMARY))
    doc.add_paragraph()
    for k, v in [("Prepared by", f("founder")),
                 ("Organisation status", f"{f('incorporation')} · {f('dpiit')}"),
                 ("Submission type", "MICE pre-application / innovation proposal"),
                 ("Phone", f("contact_phone")), ("Email", f("contact_email"))]:
        p = doc.add_paragraph()
        rr = p.add_run(f"{k}:  ")
        rr.font.bold = True
        rr.font.size = Pt(10)
        p.add_run(v).font.size = Pt(10)
    doc.add_page_break()

    for kind, payload in _content():
        if kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind in ("p", "lead"):
            p = doc.add_paragraph(_plain(payload))
            if kind == "lead":
                p.runs[0].font.size = Pt(12)
        elif kind == "ul":
            for item in payload:
                doc.add_paragraph(_plain(item), style="List Bullet")
        elif kind == "status":
            label, text = payload
            p = doc.add_paragraph()
            r = p.add_run(f"[{label}]  ")
            r.font.bold = True
            r.font.color.rgb = RGBColor(*T.rgb(T.STATUS[label]))
            p.add_run(_plain(text))
        elif kind == "table":
            rows, _ = payload
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    cl = tbl.cell(i, j)
                    cl.text = _plain(str(cell))
                    for pp in cl.paragraphs:
                        for rr in pp.runs:
                            rr.font.size = Pt(9)
                            if i == 0:
                                rr.font.bold = True
        elif kind == "callout":
            title, body, _kind = payload
            p = doc.add_paragraph()
            r = p.add_run(title + "\n")
            r.font.bold = True
            p.add_run(_plain(body)).font.size = Pt(10)
        elif kind == "flow":
            steps, _per = payload
            p = doc.add_paragraph()
            p.add_run("  →  ".join(s[0] for s in steps)).font.bold = True
        elif kind == "pagebreak":
            doc.add_page_break()
    doc.save(path)
    return path


def _plain(text: str) -> str:
    return text.replace("<b>", "").replace("</b>", "")


# -------------------------------------------------------------------- pdf
def _pdf(path: str) -> str:
    story: list = []
    for kind, payload in _content():
        if kind == "h1":
            story.append(heading(payload, 1))
        elif kind == "h2":
            story.append(heading(payload, 2))
        elif kind == "p":
            story.append(para(payload))
        elif kind == "lead":
            story.append(para(payload, "lead"))
        elif kind == "ul":
            story.extend(bullets(payload))
        elif kind == "status":
            story.append(status_row(*payload))
            story.append(Spacer(1, 1.5 * mm))
        elif kind == "table":
            rows, widths = payload
            story.append(table(rows, widths=widths))
            story.append(Spacer(1, 3 * mm))
        elif kind == "callout":
            story.append(callout(*payload))
        elif kind == "flow":
            steps, per = payload
            story.append(flow_diagram(steps, per_row=per))
        elif kind == "pagebreak":
            story.append(PageBreak())
    return build(path, "MICE Project Proposal",
                 "Commuter intelligence for the Delhi Metro",
                 "Project Proposal",
                 [("Prepared by", f("founder")),
                  ("Organisation status", f"{f('incorporation')} · {f('dpiit')}"),
                  ("Submission type", "MICE pre-application / innovation proposal"),
                  ("Phone", f("contact_phone")),
                  ("Email", f("contact_email"))],
                 story)


def build_all(docx_path: str, pdf_path: str) -> tuple[str, str]:
    return _docx(docx_path), _pdf(pdf_path)
